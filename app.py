import streamlit as st
import os
import re
import io
import zipfile
from pathlib import Path
import tempfile
from datetime import datetime
import requests
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import yt_dlp
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch, cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# ==================== PAGE CONFIG ====================
st.set_page_config(
    page_title="YT Title & Thumbnail Tool",
    page_icon="🎬",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==================== PREMIUM DARK CSS ====================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&family=Manrope:wght@600;700;800&display=swap');

:root {
    --bg: #0B0E1A;
    --card: #121629;
    --card2: #171B2E;
    --border: #252B43;
    --border-h: #323A5A;
    --primary1: #6366F1;
    --primary2: #8B5CF6;
    --blue: #3B82F6;
    --text: #F8FAFC;
    --text2: #B8C2D9;
    --muted: #94A3B8;
    --success: #22C55E;
    --warn: #F59E0B;
    --error: #EF4444;
    --r: 12px;
}

html, body, .stApp {
    background: var(--bg) !important;
    color: var(--text) !important;
    font-family: 'Inter', sans-serif !important;
}

.main .block-container {
    padding-top: 1.2rem !important;
    padding-bottom: 5rem !important;
    max-width: 1080px !important;
}

/* Headings - pure white */
h1, h2, h3, h4 {
    font-family: 'Manrope', sans-serif !important;
    color: #FFFFFF !important;
    letter-spacing: -0.02em !important;
    font-weight: 700 !important;
}

/* Body text - high contrast, but DO NOT override all div/span (breaks Streamlit) */
p, .stMarkdown p, .stCaption {
    color: var(--text2) !important;
    font-family: 'Inter', sans-serif !important;
}

/* Labels */
label, .stTextInput label, .stTextArea label, .stNumberInput label,
.stSelectbox label, .stCheckbox label, .stRadio label {
    color: #F8FAFC !important;
    font-weight: 500 !important;
    font-size: 0.85rem !important;
}

/* Sidebar */
[data-testid="stSidebar"] {
    background: #080A14 !important;
    border-right: 1px solid var(--border) !important;
}
[data-testid="stSidebar"] .stMarkdown p,
[data-testid="stSidebar"] .stCaption {
    color: #CBD5E1 !important;
}
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3,
[data-testid="stSidebar"] h4 {
    color: #FFFFFF !important;
}

/* Buttons - pure white text */
.stButton > button {
    background: linear-gradient(135deg, var(--primary1), var(--primary2)) !important;
    color: #FFFFFF !important;
    border: none !important;
    border-radius: 10px !important;
    padding: 0.5rem 1.15rem !important;
    font-weight: 600 !important;
    font-size: 0.875rem !important;
    transition: all 0.2s ease !important;
    box-shadow: 0 2px 10px rgba(99,102,241,0.28) !important;
}
.stButton > button:hover {
    transform: translateY(-1px) !important;
    box-shadow: 0 5px 18px rgba(99,102,241,0.45) !important;
    filter: brightness(1.08) !important;
}
.stButton > button p {
    color: #FFFFFF !important;
}
.stDownloadButton > button {
    background: linear-gradient(135deg, #059669, #22C55E) !important;
    color: #FFFFFF !important;
    box-shadow: 0 2px 10px rgba(34,197,94,0.25) !important;
}
.stDownloadButton > button p {
    color: #FFFFFF !important;
}

/* Inputs */
.stTextInput > div > div > input,
.stTextArea > div > div > textarea,
.stNumberInput > div > div > input {
    background: var(--card) !important;
    border: 1px solid var(--border) !important;
    border-radius: 10px !important;
    color: #F8FAFC !important;
    font-size: 0.9rem !important;
}
.stTextInput > div > div > input::placeholder,
.stTextArea > div > div > textarea::placeholder {
    color: #94A3B8 !important;
    opacity: 1 !important;
}
.stTextInput > div > div > input:focus,
.stTextArea > div > div > textarea:focus {
    border-color: var(--primary1) !important;
    box-shadow: 0 0 0 3px rgba(99,102,241,0.18) !important;
}
.stSelectbox > div > div {
    background: var(--card) !important;
    border: 1px solid var(--border) !important;
    border-radius: 10px !important;
    color: #F8FAFC !important;
}

/* Expander header text - high contrast */
[data-testid="stExpander"] summary,
[data-testid="stExpander"] summary p,
[data-testid="stExpander"] summary span {
    color: #F8FAFC !important;
    font-weight: 500 !important;
}

/* Tabs */
.stTabs [data-baseweb="tab-list"] {
    gap: 4px; background: var(--card); border-radius: 12px;
    padding: 4px; border: 1px solid var(--border);
}
.stTabs [data-baseweb="tab"] {
    border-radius: 9px !important;
    color: #CBD5E1 !important;
    font-weight: 500 !important;
    font-size: 0.875rem !important;
}
.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, var(--primary1), var(--primary2)) !important;
    color: #FFFFFF !important;
    box-shadow: 0 0 16px rgba(99,102,241,0.35) !important;
}

/* File uploader */
[data-testid="stFileUploader"] {
    background: var(--card) !important;
    border: 1.5px dashed var(--border-h) !important;
    border-radius: 12px !important;
    padding: 1rem !important;
}

/* Metric */
[data-testid="stMetric"] {
    background: var(--card) !important;
    border: 1px solid var(--border) !important;
    border-radius: 10px !important;
    padding: 0.75rem !important;
}
[data-testid="stMetric"] label {
    color: #CBD5E1 !important;
}
[data-testid="stMetric"] [data-testid="stMetricValue"] {
    color: #FFFFFF !important;
}

/* Alerts */
.stSuccess { background: rgba(34,197,94,0.1) !important; border: 1px solid rgba(34,197,94,0.25) !important; border-radius: 10px !important; }
.stError { background: rgba(239,68,68,0.1) !important; border: 1px solid rgba(239,68,68,0.25) !important; border-radius: 10px !important; }
.stWarning, .stInfo { background: rgba(245,158,11,0.08) !important; border: 1px solid rgba(245,158,11,0.2) !important; border-radius: 10px !important; }

/* Progress */
.stProgress > div > div {
    background: linear-gradient(90deg, var(--primary1), var(--primary2)) !important;
    border-radius: 6px !important;
}

hr { border-color: var(--border) !important; margin: 0.6rem 0 !important; }
/* Keep sidebar open/close control visible */
#MainMenu { visibility: hidden; }
footer { visibility: hidden; }
header { visibility: visible !important; }
[data-testid="stSidebarCollapsedControl"],
[data-testid="collapsedControl"],
section[data-testid="stSidebar"] + div button,
button[kind="header"] {
    visibility: visible !important;
    display: flex !important;
    opacity: 1 !important;
    color: #F8FAFC !important;
    z-index: 999999 !important;
}

::-webkit-scrollbar { width: 6px; }
::-webkit-scrollbar-track { background: var(--bg); }
::-webkit-scrollbar-thumb { background: var(--border-h); border-radius: 4px; }

[data-testid="stImage"] img {
    border-radius: 10px !important;
    border: 1px solid var(--border) !important;
}

.badge {
    display: inline-block;
    background: rgba(99,102,241,0.2);
    color: #E0E7FF !important;
    border: 1px solid rgba(99,102,241,0.4);
    font-size: 0.75rem;
    font-weight: 600;
    padding: 3px 10px;
    border-radius: 20px;
}
.num-badge {
    display: inline-flex; align-items: center; justify-content: center;
    background: linear-gradient(135deg, var(--primary1), var(--primary2));
    color: #FFFFFF !important; font-weight: 700; font-family: 'Manrope', sans-serif;
    padding: 2px 10px; border-radius: 20px; font-size: 0.78rem;
    box-shadow: 0 2px 8px rgba(99,102,241,0.3);
}
.step-active { color: #FFFFFF !important; font-weight: 600 !important; }
.step-done { color: var(--success) !important; }
.step-pending { color: #94A3B8 !important; }

.stCaption, [data-testid="stCaption"] {
    color: #B8C2D9 !important;
}
</style>
""", unsafe_allow_html=True)

# ==================== HELPERS ====================

def is_cloud_host():
    """This build is Streamlit Cloud only — always cloud UI."""
    return True


def extract_video_id(url: str):
    patterns = [
        r'(?:youtube\.com\/watch\?v=|youtu\.be\/|youtube\.com\/embed\/|youtube\.com\/v\/)([a-zA-Z0-9_-]{11})',
        r'^([a-zA-Z0-9_-]{11})$'
    ]
    for p in patterns:
        m = re.search(p, url.strip())
        if m:
            return m.group(1)
    return None

def is_shorts_url(url: str) -> bool:
    return "youtube.com/shorts/" in url.lower() or "/shorts/" in url.lower()

def get_thumb_url(vid: str) -> str:
    return f"https://i.ytimg.com/vi/{vid}/maxresdefault.jpg"

def fetch_info(url: str) -> dict:
    ydl_opts = {'quiet': True, 'no_warnings': True, 'skip_download': True}
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
            return {
                'ok': True,
                'title': info.get('title', 'Unknown'),
                'video_id': info.get('id'),
                'url': url,
                'channel': info.get('uploader', ''),
            }
    except Exception as e:
        vid = extract_video_id(url)
        if vid:
            return {'ok': True, 'title': f"Video_{vid}", 'video_id': vid, 'url': url, 'channel': '', 'warn': str(e)}
        return {'ok': False, 'error': str(e)}


def normalize_channel_url(url: str) -> str:
    """Prefer /videos tab for channel listing."""
    u = url.strip().rstrip("/")
    if not u:
        return u
    low = u.lower()
    if "/videos" in low or "/streams" in low or "/shorts" in low or "/playlists" in low:
        return u
    # @handle or /channel/ /c/ /user/
    if any(x in low for x in ["/@", "/channel/", "/c/", "/user/"]):
        return u + "/videos"
    return u

def fetch_channel_videos(channel_url: str, max_videos: int = 50) -> dict:
    """
    List long-form videos from a channel (titles + ids only, no file download).
    Returns {ok, items:[{title, video_id, url}], error?}
    """
    url = normalize_channel_url(channel_url)
    ydl_opts = {
        "quiet": True,
        "no_warnings": True,
        "skip_download": True,
        "extract_flat": "in_playlist",
        "playlistend": max(1, min(int(max_videos), 500)),
        "ignoreerrors": True,
    }
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
        if not info:
            return {"ok": False, "error": "Could not fetch channel data", "items": []}
        entries = info.get("entries") or []
        items = []
        seen = set()
        for e in entries:
            if not e:
                continue
            vid = e.get("id") or extract_video_id(e.get("url") or e.get("webpage_url") or "")
            title = (e.get("title") or "").strip()
            if not vid or vid in seen:
                continue
            # skip obvious shorts by url/title heuristics
            webpage = (e.get("url") or e.get("webpage_url") or "") or ""
            if "shorts" in webpage.lower() or (e.get("duration") is not None and e.get("duration") and e.get("duration") < 60):
                # duration may be missing in flat extract — still skip shorts path
                if "shorts" in webpage.lower():
                    continue
            if not title or title == "[Deleted video]" or title == "[Private video]":
                continue
            seen.add(vid)
            items.append({
                "title": title,
                "video_id": vid,
                "url": f"https://www.youtube.com/watch?v={vid}",
            })
            if len(items) >= max_videos:
                break
        return {"ok": True, "items": items, "channel": info.get("channel") or info.get("uploader") or info.get("title") or ""}
    except Exception as e:
        return {"ok": False, "error": str(e), "items": []}

def download_thumb(vid: str, path: str) -> bool:
    """Download thumbnail to local path. Returns True on success."""
    for q in ["maxresdefault", "sddefault", "hqdefault", "mqdefault"]:
        try:
            r = requests.get(f"https://i.ytimg.com/vi/{vid}/{q}.jpg", timeout=12)
            if r.status_code == 200 and len(r.content) > 1500:
                with open(path, 'wb') as f:
                    f.write(r.content)
                return True
        except:
            continue
    return False

def get_thumb_bytes(vid: str):
    """Fetch thumbnail bytes without writing to disk. Returns bytes or None."""
    for q in ["maxresdefault", "sddefault", "hqdefault", "mqdefault"]:
        try:
            r = requests.get(f"https://i.ytimg.com/vi/{vid}/{q}.jpg", timeout=12)
            if r.status_code == 200 and len(r.content) > 1500:
                return r.content
        except:
            continue
    return None


def browse_folder():
    """Open Windows folder picker. Returns selected path or None."""
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.withdraw()
        root.attributes("-topmost", True)
        path = filedialog.askdirectory(title="Select folder to save thumbnails")
        root.destroy()
        return path if path else None
    except Exception as e:
        return None


def require_save_folder():
    """If user has not chosen a folder, open Browse dialog. Returns True if folder is ready."""
    if st.session_state.get("folder_chosen"):
        Path(st.session_state.output_dir).mkdir(parents=True, exist_ok=True)
        return True
    chosen = browse_folder()
    if not chosen:
        st.warning("Please select a folder first.")
        return False
    p = Path(chosen).resolve()
    p.mkdir(parents=True, exist_ok=True)
    st.session_state.output_dir = p
    st.session_state.folder_chosen = True
    return True

def ensure_local_thumb(it, num_fmt, auto_num):
    """Ensure thumbnail is saved to user-chosen output folder. Returns local path or None."""
    out = Path(st.session_state.output_dir)
    out.mkdir(parents=True, exist_ok=True)
    # If already saved inside current output folder, reuse
    if it.get('thumb_path') and os.path.exists(it['thumb_path']):
        return it['thumb_path']
    vid = it.get('video_id')
    if not vid:
        return None
    prefix = format_number(it.get('number') or 1, num_fmt) if auto_num else ""
    fname = f"{prefix}{safe_name(it.get('my_title') or it.get('competitor_title') or vid)}.jpg"
    fpath = out / fname
    if download_thumb(vid, str(fpath)):
        it['thumb_path'] = str(fpath)
        return str(fpath)
    return None

def _para_yt_url(para):
    """Return long-form YouTube URL from paragraph hyperlinks or plain text, else None."""
    urls = []
    try:
        for h in para.hyperlinks:
            if h.url:
                urls.append(h.url)
    except Exception:
        pass
    text = (para.text or "").strip()
    # also catch plain URLs in text
    for m in re.finditer(r'https?://[^\s<>"\']+', text):
        urls.append(m.group(0))
    for u in urls:
        if is_shorts_url(u):
            continue
        if extract_video_id(u):
            return u
    # plain video id in text
    vid = extract_video_id(text)
    if vid:
        return f"https://www.youtube.com/watch?v={vid}"
    return None

def extract_docx_pairs(file_bytes) -> list:
    """
    Word format expected:
      Line 1: Competitor title + YouTube hyperlink
      Line 2: My Title (optional; may contain non-YT links — ignored)
    Only long-form YouTube links are kept.
    Returns list of {competitor_title, my_title, url}
    """
    doc = Document(io.BytesIO(file_bytes))
    paras = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if not t:
            continue
        paras.append(para)
    # tables too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = (para.text or "").strip()
                    if not t:
                        continue
                    paras.append(para)

    pairs = []
    i = 0
    while i < len(paras):
        para = paras[i]
        yt = _para_yt_url(para)
        if yt:
            competitor = (para.text or "").strip()
            my_title = ""
            # next line = My Title if it does NOT contain a YT long-form link
            if i + 1 < len(paras):
                nxt = paras[i + 1]
                if not _para_yt_url(nxt):
                    my_title = (nxt.text or "").strip()
                    i += 1  # consume my-title line
            pairs.append({
                "competitor_title": competitor,
                "my_title": my_title,
                "url": yt,
            })
        i += 1
    return pairs

def parse_paste_block(text: str) -> list:
    """
    Parse pasted text. Supports plain YT links, labeled blocks, title+URL.
    Returns list of {competitor_title, my_title, url}
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    items = []

    def is_yt(s):
        return bool(extract_video_id(s)) and not is_shorts_url(s)

    def to_url(s):
        vid = extract_video_id(s)
        if not vid:
            return None
        if s.startswith("http"):
            return s.split()[0]
        return f"https://www.youtube.com/watch?v={vid}"

    def strip_label(line, *labels):
        low = line.lower().strip()
        for lab in labels:
            if low.startswith(lab):
                rest = line[len(lab):]
                rest = rest.lstrip(" \t:-–—")
                return rest.strip()
        return None

    i = 0
    while i < len(lines):
        line = lines[i]
        low = line.lower()

        comp_rest = strip_label(line, "competitor title", "competitor", "yt title", "youtube title")
        if comp_rest is not None or low.startswith("competitor"):
            comp = comp_rest if comp_rest is not None else (line.split(":", 1)[1].strip() if ":" in line else "")
            my_t = ""
            url = None
            j = i + 1
            while j < len(lines):
                lj = lines[j]
                lj_low = lj.lower()
                my_rest = strip_label(lj, "my title", "my_title", "mytitle")
                if my_rest is not None or lj_low.startswith("my title"):
                    my_t = my_rest if my_rest is not None else (lj.split(":", 1)[1].strip() if ":" in lj else "")
                    j += 1
                    continue
                if is_yt(lj):
                    url = to_url(lj)
                    j += 1
                    break
                if strip_label(lj, "competitor title", "competitor") is not None or lj_low.startswith("competitor"):
                    break
                j += 1
            if url:
                items.append({"competitor_title": comp, "my_title": my_t, "url": url})
            i = max(j, i + 1)
            continue

        if is_yt(line):
            items.append({"competitor_title": "", "my_title": "", "url": to_url(line)})
            i += 1
            continue

        url_idx = None
        for k in range(i + 1, min(i + 6, len(lines))):
            if is_yt(lines[k]):
                url_idx = k
                break
            if strip_label(lines[k], "competitor title", "competitor") is not None:
                break
        if url_idx is not None:
            comp = line
            my_t = ""
            for k in range(i + 1, url_idx):
                my_rest = strip_label(lines[k], "my title", "my_title", "mytitle")
                if my_rest is not None or lines[k].lower().startswith("my title"):
                    my_t = my_rest if my_rest is not None else (
                        lines[k].split(":", 1)[1].strip() if ":" in lines[k] else lines[k]
                    )
                elif not is_yt(lines[k]) and not my_t:
                    my_t = lines[k]
            items.append({"competitor_title": comp, "my_title": my_t, "url": to_url(lines[url_idx])})
            i = url_idx + 1
            continue

        i += 1
    return items


def safe_name(name: str, maxlen=70) -> str:
    name = re.sub(r'[<>:"/\\|?*]', '', name)
    name = re.sub(r'\s+', ' ', name).strip()
    return name[:maxlen]

def format_number(n: int, fmt: str) -> str:
    if fmt == "1. Title":
        return f"{n}. "
    if fmt == "01. Title":
        return f"{n:02d}. "
    if fmt == "001. Title":
        return f"{n:03d}. "
    if fmt == "1 - Title":
        return f"{n} - "
    if fmt == "01 - Title":
        return f"{n:02d} - "
    return f"{n:03d} - "

def make_pdf(items, path, num_fmt):
    """Build PDF with embedded thumbs. Thumbs use a TEMP folder only — user folder gets ONLY the PDF."""
    doc = SimpleDocTemplate(path, pagesize=A4,
                            rightMargin=1.4*cm, leftMargin=1.4*cm,
                            topMargin=1.4*cm, bottomMargin=1.4*cm)
    styles = getSampleStyleSheet()
    s_title = ParagraphStyle('T', parent=styles['Heading2'], fontSize=12, spaceAfter=4, fontName='Helvetica-Bold')
    s_norm = ParagraphStyle('N', parent=styles['Normal'], fontSize=10, spaceAfter=3)
    s_num = ParagraphStyle('Num', parent=styles['Normal'], fontSize=11, textColor='#6366F1', fontName='Helvetica-Bold', spaceAfter=6)
    story = []
    with tempfile.TemporaryDirectory(prefix="yt_pdf_thumbs_") as tmp:
        tmp_path = Path(tmp)
        for it in items:
            els = []
            prefix = format_number(it['number'], num_fmt)
            els.append(Paragraph(f"<b>#{it['number']:03d}</b>", s_num))
            # Prefer already-saved user thumb if exists; else download to TEMP only
            img_file = None
            if it.get('thumb_path') and os.path.exists(it['thumb_path']):
                img_file = it['thumb_path']
            elif it.get('video_id'):
                tfile = tmp_path / f"{it['video_id']}.jpg"
                if download_thumb(it['video_id'], str(tfile)):
                    img_file = str(tfile)
            if img_file:
                try:
                    els.append(RLImage(img_file, width=4.2*inch, height=2.36*inch))
                    els.append(Spacer(1, 6))
                except Exception:
                    pass
            els.append(Paragraph(f"<b>Competitor Title:</b> {it.get('competitor_title','')}", s_norm))
            my = it.get('my_title') or ""
            if my:
                els.append(Paragraph(f"<b>My Title:</b> {prefix}{my}", s_title))
            else:
                els.append(Paragraph(f"<b>My Title:</b> (blank)", s_norm))
            els.append(Spacer(1, 14))
            story.append(KeepTogether(els))
        doc.build(story)
    # temp folder auto-deleted — only PDF remains at `path`

# ==================== SESSION ====================
if 'video_items' not in st.session_state:
    st.session_state.video_items = []
if 'page' not in st.session_state:
    st.session_state.page = "add"
if 'output_dir' not in st.session_state:
    st.session_state.output_dir = Path("output").resolve()
if 'folder_chosen' not in st.session_state:
    st.session_state.folder_chosen = False  # True only after Browse/Apply
if 'last_saved' not in st.session_state:
    st.session_state.last_saved = []
if 'word_scan_results' not in st.session_state:
    st.session_state.word_scan_results = []
if 'paste_scan_results' not in st.session_state:
    st.session_state.paste_scan_results = []

# ==================== SIDEBAR ====================
with st.sidebar:
    st.markdown("### 🎬 YT Title & Thumbnail Tool")
    st.caption("Long-form YouTube videos only")
    st.markdown("---")

    st.markdown("**MAIN**")
    if st.button("Add Links", use_container_width=True,
                 type="primary" if st.session_state.page == "add" else "secondary"):
        st.session_state.page = "add"
        st.rerun()
    if st.button("Review & Export", use_container_width=True,
                 type="primary" if st.session_state.page == "review" else "secondary"):
        st.session_state.page = "review"
        st.rerun()

    st.markdown("---")
    st.markdown("**SETTINGS**")

    with st.expander("Save Location", expanded=True):
        st.caption("Output folder")
        st.caption(f"Current: `{st.session_state.output_dir}`")
        c_browse, c_apply = st.columns(2)
        with c_browse:
            if st.button("Browse Folder", use_container_width=True, key="browse_folder_btn"):
                chosen = browse_folder()
                if chosen:
                    p = Path(chosen).resolve()
                    p.mkdir(parents=True, exist_ok=True)
                    st.session_state.output_dir = p
                    st.session_state.folder_chosen = True
                    st.session_state["out_path_input"] = str(p)
                    st.success(f"Folder set:\n{p}")
                    st.rerun()
                else:
                    st.info("No folder selected")
        with c_apply:
            pass
        new_path = st.text_input(
            "Or type path",
            value=str(st.session_state.output_dir),
            key="out_path_input"
        )
        if st.button("Apply Typed Path", use_container_width=True, key="apply_folder"):
            try:
                p = Path(new_path.strip()).expanduser().resolve()
                p.mkdir(parents=True, exist_ok=True)
                st.session_state.output_dir = p
                st.session_state.folder_chosen = True
                st.success(f"Folder set:\n{p}")
            except Exception as e:
                st.error(f"Invalid path: {e}")

    with st.expander("Thumbnail Settings", expanded=False):
        st.caption("Quality: HD 1280x720 (maxres)")
        st.caption("Fallback to lower quality if needed")

    with st.expander("Title & Numbering", expanded=False):
        auto_num = st.checkbox("Auto Number Titles", value=True)
        start_num = st.number_input("Start Number", min_value=1, value=1, step=1)
        num_fmt = st.selectbox("Number Format", [
            "001. Title", "01. Title", "1. Title", "01 - Title", "1 - Title"
        ], index=0)

    st.markdown("---")
    st.markdown("**CURRENT SESSION**")
    st.metric("Items Loaded", len(st.session_state.video_items))
    if st.button("Clear All", use_container_width=True):
        st.session_state.video_items = []
        st.session_state.last_saved = []
        st.rerun()

# ==================== HEADER ====================
st.markdown("<h1 style='margin-bottom:2px;'>🎬 YT Title & Thumbnail Tool</h1>", unsafe_allow_html=True)
st.markdown("<p style='color:#B8C2D9; margin-top:0; font-size:0.95rem;'>Extract titles • Download thumbnails • Generate numbered output</p>", unsafe_allow_html=True)

# Workflow steps
c1, c2, c3 = st.columns(3)
with c1:
    cls = "step-active" if st.session_state.page == "add" else ("step-done" if st.session_state.video_items else "step-pending")
    st.markdown(f"<p class='{cls}'>① Add Links<br><span style='font-size:0.8rem;font-weight:400;'>Add your YouTube videos</span></p>", unsafe_allow_html=True)
with c2:
    cls = "step-active" if st.session_state.page == "review" else ("step-done" if st.session_state.video_items else "step-pending")
    st.markdown(f"<p class='{cls}'>② Review & Edit<br><span style='font-size:0.8rem;font-weight:400;'>Review titles & thumbnails</span></p>", unsafe_allow_html=True)
with c3:
    cls = "step-pending"
    st.markdown(f"<p class='{cls}'>③ Export<br><span style='font-size:0.8rem;font-weight:400;'>Download your content</span></p>", unsafe_allow_html=True)

st.markdown("")

# ==================== PAGE: ADD LINKS ====================
if st.session_state.page == "add":

    col_main, col_side = st.columns([2.2, 1])

    with col_main:
        st.markdown("### Add YouTube Links")
        st.markdown('<span class="badge">✓ Long-form videos only</span>', unsafe_allow_html=True)
        st.write("")

        st.markdown("**Single YouTube Link**")
        single = st.text_input("single_link", placeholder="Paste a YouTube video link here...", label_visibility="collapsed")
        add_one = st.button("Add Link", use_container_width=True)

        st.markdown("")
        st.markdown("**YouTube Channel**")
        st.caption("Paste a channel URL to list long-form video titles (thumbnails on Review).")
        ch_url = st.text_input(
            "channel_url",
            placeholder="https://www.youtube.com/@ChannelName  or  /channel/UCxxxx",
            label_visibility="collapsed",
        )
        c1, c2 = st.columns([1, 1])
        with c1:
            ch_max = st.number_input("Max videos", min_value=1, max_value=300, value=50, step=10, key="ch_max")
        with c2:
            st.write("")
            st.write("")
            fetch_ch = st.button("Fetch Channel Videos", use_container_width=True, key="fetch_channel")

        if fetch_ch and ch_url.strip():
            with st.spinner("Channel se video list nikal rahe hain..."):
                result = fetch_channel_videos(ch_url.strip(), max_videos=int(ch_max))
            if not result.get("ok"):
                st.error(f"Channel fetch fail: {result.get('error', 'unknown')}")
            else:
                items = result.get("items") or []
                existing_ids = {it.get("video_id") for it in st.session_state.video_items if it.get("video_id")}
                added = 0
                skipped = 0
                for row in items:
                    vid = row.get("video_id")
                    if not vid or vid in existing_ids:
                        skipped += 1
                        continue
                    st.session_state.video_items.append({
                        "competitor_title": row.get("title") or "",
                        "my_title": "",  # blank — user fills
                        "url": row.get("url"),
                        "video_id": vid,
                        "number": None,
                        "thumb_path": None,
                    })
                    existing_ids.add(vid)
                    added += 1
                ch_name = result.get("channel") or ""
                st.success(
                    f"Channel{(' — ' + ch_name) if ch_name else ''}: "
                    f"{len(items)} found, {added} added"
                    + (f", {skipped} skipped (duplicate)" if skipped else "")
                )
                if added > 0:
                    st.session_state.page = "review"
                    st.rerun()

        st.markdown("")
        st.markdown("**Bulk / Paste**")
        st.caption("Plain YT links, OR title+URL lines, OR labeled: Competitor Title: / My Title:")
        bulk = st.text_area(
            "bulk_links",
            height=160,
            placeholder=(
                "Option A — links only:\n"
                "https://www.youtube.com/watch?v=...\n"
                "https://youtu.be/...\n\n"
                "Option B — with titles:\n"
                "Competitor Title: Their video title\n"
                "My Title: Your version title\n"
                "https://www.youtube.com/watch?v=...\n"
            ),
            label_visibility="collapsed",
        )

        cscan, cproc = st.columns(2)
        with cscan:
            scan_paste = st.button("Scan Paste", use_container_width=True)
        with cproc:
            process_btn = st.button("Process Links", use_container_width=True, type="primary")

        if scan_paste and bulk.strip():
            parsed = parse_paste_block(bulk)
            parsed = [x for x in parsed if x.get("url") and not is_shorts_url(x["url"]) and extract_video_id(x["url"])]
            st.session_state.paste_scan_results = parsed
            if parsed:
                st.success(f"Scan Complete — {len(parsed)} long-form YouTube link(s) found.")
                with st.expander("Preview scanned items", expanded=True):
                    for i, it in enumerate(parsed, 1):
                        st.caption(
                            f"{i}. Comp: {it.get('competitor_title') or '(blank)'} | "
                            f"My: {it.get('my_title') or '(blank)'} | {it.get('url','')[:60]}"
                        )
            else:
                st.error(
                    "No YouTube URL found.\n\n"
                    "Copying titles from Google Docs often drops the hyperlink (text only).\n\n"
                    "**Working options:**\n"
                    "1) Upload a **Word .docx** (hyperlinks are preserved) — recommended\n"
                    "2) Paste in this format with the URL included:\n"
                    "```\n"
                    "Competitor title text here\n"
                    "My Title: your title here\n"
                    "https://www.youtube.com/watch?v=VIDEO_ID\n"
                    "```"
                )

        if st.session_state.get("paste_scan_results"):
            st.info(f"Scanned paste ready: {len(st.session_state.paste_scan_results)} item(s). Click Process Links to add.")

        st.markdown("")
        with st.expander("Upload Word file (.docx)", expanded=True):
            st.caption("Format: Line1 = Competitor title + YT hyperlink · Line2 = My Title (optional)")
            up = st.file_uploader("Upload .docx", type=["docx"], label_visibility="collapsed", key="docx_up")
            if up is not None:
                if st.button("Scan & Add Word File", use_container_width=True, type="primary", key="scan_add_word"):
                    # Prevent double-run on same upload
                    file_id = f"{up.name}_{up.size}"
                    if st.session_state.get("last_word_file_id") == file_id and st.session_state.get("word_just_added"):
                        st.warning("This file was already added. Upload a new file or click Clear All.")
                    else:
                        with st.spinner("Scanning Word + adding (one time)..."):
                            data = up.read()
                            pairs = extract_docx_pairs(data)
                            # existing URLs/ids to skip duplicates
                            existing_ids = {it.get("video_id") for it in st.session_state.video_items if it.get("video_id")}
                            existing_urls = {it.get("url") for it in st.session_state.video_items if it.get("url")}
                            added = 0
                            skipped = 0
                            bar = st.progress(0)
                            status = st.empty()
                            for i, row in enumerate(pairs):
                                url = row.get("url")
                                status.text(f"Adding {i+1} of {len(pairs)}...")
                                if not url or is_shorts_url(url):
                                    skipped += 1
                                    bar.progress((i + 1) / max(len(pairs), 1))
                                    continue
                                vid = extract_video_id(url)
                                if vid in existing_ids or url in existing_urls:
                                    skipped += 1
                                    bar.progress((i + 1) / max(len(pairs), 1))
                                    continue
                                info = fetch_info(url)
                                if info.get("ok"):
                                    vid = info.get("video_id") or vid
                                    url = info.get("url") or url
                                if vid in existing_ids:
                                    skipped += 1
                                    bar.progress((i + 1) / max(len(pairs), 1))
                                    continue
                                st.session_state.video_items.append({
                                    "competitor_title": row.get("competitor_title") or "",
                                    "my_title": row.get("my_title") or "",
                                    "url": url,
                                    "video_id": vid,
                                    "number": None,
                                    "thumb_path": None,
                                })
                                if vid:
                                    existing_ids.add(vid)
                                if url:
                                    existing_urls.add(url)
                                added += 1
                                bar.progress((i + 1) / max(len(pairs), 1))
                            status.empty()
                            bar.empty()
                            st.session_state.word_scan_results = []
                            st.session_state.last_word_file_id = file_id
                            st.session_state.word_just_added = True
                            st.success(
                                f"Done — scanned {len(pairs)}, added {added}"
                                + (f", skipped {skipped} duplicate/invalid" if skipped else "")
                                + "."
                            )
                            if added > 0:
                                st.session_state.page = "review"
                                st.rerun()

    with col_side:
        st.markdown("### 💡 How it works")
        st.markdown("""
1. Add YouTube video links  
2. Titles & thumbnails are extracted  
3. Review, edit and export  
        """)
        st.info("⭐ Tip: Make sure the videos are public so their data can be fetched successfully.")
        st.markdown("---")
        st.markdown("**Quick Features**")
        st.caption("• Extract Titles")
        st.caption("• HD Thumbnails (1280×720)")
        st.caption("• Auto Numbering")
        st.caption("• Bulk Export")

    if add_one and single.strip():
        url = single.strip()
        if is_shorts_url(url):
            st.warning("This tool is for long-form YouTube videos only.")
        elif not extract_video_id(url):
            st.error("Invalid YouTube URL")
        else:
            with st.spinner("Fetching..."):
                info = fetch_info(url)
                if info.get("ok"):
                    st.session_state.video_items.append({
                        "competitor_title": info["title"],
                        "my_title": "",  # blank — user fills manually
                        "url": info["url"],
                        "video_id": info["video_id"],
                        "number": None,
                        "thumb_path": None,
                    })
                    st.success(f"Added: {info['title'][:55]}... (My Title left blank)")
                    st.session_state.page = "review"
                    st.rerun()
                else:
                    st.error("Could not process this video")

    if process_btn:
        to_add = []
        # prefer scanned paste results if available
        if st.session_state.get("paste_scan_results"):
            to_add = list(st.session_state.paste_scan_results)
        else:
            if single.strip():
                to_add.append({
                    "competitor_title": "",
                    "my_title": "",
                    "url": single.strip(),
                })
            if bulk.strip():
                parsed = parse_paste_block(bulk)
                if parsed:
                    to_add.extend(parsed)
                else:
                    for l in bulk.strip().splitlines():
                        l = l.strip()
                        if l and extract_video_id(l) and not is_shorts_url(l):
                            to_add.append({"competitor_title": "", "my_title": "", "url": l})

        if not to_add:
            st.warning("Please add at least one valid long-form YouTube link.")
        else:
            bar = st.progress(0)
            status = st.empty()
            added = 0
            for i, row in enumerate(to_add):
                u = row.get("url") or ""
                status.text(f"Processing {i+1} of {len(to_add)}...")
                if is_shorts_url(u) or not extract_video_id(u):
                    bar.progress((i+1)/len(to_add))
                    continue
                info = fetch_info(u)
                if info.get("ok"):
                    # Manual/paste: competitor = given or YT title; my_title only if user provided
                    comp = (row.get("competitor_title") or "").strip() or info["title"]
                    my_t = (row.get("my_title") or "").strip()  # blank unless provided
                    st.session_state.video_items.append({
                        "competitor_title": comp,
                        "my_title": my_t,
                        "url": info["url"],
                        "video_id": info["video_id"],
                        "number": None,
                        "thumb_path": None,
                    })
                    added += 1
                bar.progress((i+1)/len(to_add))
            status.empty()
            bar.empty()
            st.session_state.paste_scan_results = []
            st.success(f"{added} videos processed. My Title left blank where not provided.")
            st.session_state.page = "review"
            st.rerun()


# ==================== PAGE: REVIEW & EXPORT ====================
elif st.session_state.page == "review":

    items = st.session_state.video_items

    if not items:
        st.markdown("### Review & Export")
        st.info("No videos added yet\n\nPaste your YouTube links above to get started.")
        if st.button("← Go to Add Links"):
            st.session_state.page = "add"
            st.rerun()
    else:
        for idx, it in enumerate(items):
            it['number'] = start_num + idx

        st.markdown(f"### Review & Export  ·  {len(items)} videos")

        search = st.text_input("search", placeholder="🔍 Search titles...", label_visibility="collapsed")
        filtered = items
        if search.strip():
            q = search.strip().lower()
            filtered = [it for it in items if q in (it.get('my_title') or '').lower() or q in (it.get('competitor_title') or '').lower()]

        b1, b2, b3, b4 = st.columns(4)
        with b1:
            if st.button("Select All", use_container_width=True):
                for it in items:
                    it['selected'] = True
                st.rerun()
        with b2:
            if st.button("Deselect All", use_container_width=True):
                for it in items:
                    it['selected'] = False
                st.rerun()
        with b3:
            if st.button("Delete Selected", use_container_width=True):
                st.session_state.video_items = [it for it in items if not it.get('selected')]
                st.rerun()
        with b4:
            dl_sel = st.button("Download Selected Thumbs", use_container_width=True)

        st.markdown("---")

        for idx, it in enumerate(filtered):
            real_idx = items.index(it)
            cols = st.columns([0.08, 0.28, 0.48, 0.16])

            with cols[0]:
                st.markdown(f'<span class="num-badge">{it["number"]:02d}</span>', unsafe_allow_html=True)
                it['selected'] = st.checkbox("sel", value=it.get('selected', False), key=f"sel_{real_idx}", label_visibility="collapsed")

            with cols[1]:
                # Preview from memory (no remote URL, no auto disk save → )
                shown = False
                if it.get('thumb_path') and os.path.exists(it['thumb_path']):
                    st.image(it['thumb_path'], use_container_width=True)
                    shown = True
                elif it.get('video_id'):
                    data = get_thumb_bytes(it['video_id'])
                    if data:
                        st.image(data, use_container_width=True)
                        shown = True
                if not shown:
                    st.caption("No thumbnail")

            with cols[2]:
                st.caption("Original Title")
                st.text(it.get('competitor_title', '')[:90])
                new_title = st.text_input(
                    "My Title",
                    value=it.get('my_title') or "",
                    key=f"title_{real_idx}"
                )
                it['my_title'] = new_title
                if it.get('url'):
                    st.caption(it['url'][:65] + "...")

            with cols[3]:
                if it.get('video_id'):
                    # Direct browser download — no folder select (local + Streamlit)
                    img_bytes = get_thumb_bytes(it['video_id'])
                    if img_bytes:
                        prefix = format_number(it['number'], num_fmt) if auto_num else ""
                        base = safe_name(it.get('my_title') or it.get('competitor_title') or it['video_id'])
                        fname = f"{prefix}{base}.jpg"
                        st.download_button(
                            "Download HD",
                            data=img_bytes,
                            file_name=fname,
                            mime="image/jpeg",
                            key=f"dl_{real_idx}",
                            use_container_width=True,
                        )
                    else:
                        st.caption("Thumb N/A")
                if st.button("Remove", key=f"rm_{real_idx}", use_container_width=True):
                    st.session_state.video_items.pop(real_idx)
                    st.rerun()

            st.markdown("---")

        if dl_sel:
            selected = [it for it in items if it.get('selected') and it.get('video_id')]
            if not selected:
                st.warning("No items selected")
            else:
                # Always prepare ZIP of selected for direct download (no folder needed)
                zbuf = io.BytesIO()
                n = 0
                with zipfile.ZipFile(zbuf, "w", zipfile.ZIP_DEFLATED) as zf:
                    for it in selected:
                        b = get_thumb_bytes(it["video_id"])
                        if not b:
                            continue
                        prefix = format_number(it["number"], num_fmt) if auto_num else ""
                        base = safe_name(it.get("my_title") or it.get("competitor_title") or it["video_id"])
                        zf.writestr(f"{prefix}{base}.jpg", b)
                        n += 1
                zbuf.seek(0)
                if n:
                    st.session_state["selected_zip"] = zbuf.getvalue()
                    st.success(f"{n} selected thumbs ready — download below")
                else:
                    st.warning("Could not fetch selected thumbnails")
            if st.session_state.get("selected_zip"):
                st.download_button(
                    "Download Selected ZIP",
                    data=st.session_state["selected_zip"],
                    file_name="selected_thumbs.zip",
                    mime="application/zip",
                    use_container_width=True,
                    key="dl_selected_zip",
                )

        # Show last saved files
        if st.session_state.get('last_saved'):
            with st.expander(f"Last saved ({len(st.session_state.last_saved)} file(s))", expanded=True):
                for p in st.session_state.last_saved[-20:]:
                    st.code(p, language=None)
                st.caption(f"Folder: {st.session_state.output_dir}")


        st.markdown("### Export")
        st.caption("Download files with the buttons below.")

        e1, e2, e3 = st.columns(3)

        # --- Titles ---
        txt_lines = []
        for it in items:
            prefix = format_number(it['number'], num_fmt) if auto_num else ""
            txt_lines.append(f"{prefix}{it.get('my_title') or it.get('competitor_title','')}")
        txt_data = "\n".join(txt_lines)
        csv_lines = ["Number,Competitor Title,My Title,URL"]
        for it in items:
            csv_lines.append(
                f'{it["number"]},"{it.get("competitor_title","")}","{it.get("my_title","")}","{it.get("url","")}"'
            )
        csv_data = "\n".join(csv_lines)

        with e1:
            st.markdown("**Titles**")
            st.download_button("Download TXT", txt_data, file_name="titles.txt",
                               mime="text/plain", use_container_width=True, key="dl_txt_browser")
            st.download_button("Download CSV", csv_data, file_name="titles.csv",
                               mime="text/csv", use_container_width=True, key="dl_csv_browser")
            if not is_cloud_host():
                if st.button("Save TXT to Folder", use_container_width=True, key="save_txt"):
                    if require_save_folder():
                        tp = Path(st.session_state.output_dir) / "titles.txt"
                        tp.write_text(txt_data, encoding="utf-8")
                        st.session_state.last_saved = [str(tp)]
                        st.success(f"Saved: {tp}")
                if st.button("Save CSV to Folder", use_container_width=True, key="save_csv"):
                    if require_save_folder():
                        cp = Path(st.session_state.output_dir) / "titles.csv"
                        cp.write_text(csv_data, encoding="utf-8")
                        st.session_state.last_saved = [str(cp)]
                        st.success(f"Saved: {cp}")

        with e2:
            st.markdown("**Thumbnails**")
            st.caption("Per item: Download HD. Bulk ZIP below.")
            # ZIP in memory for browser download
            zip_buf = io.BytesIO()
            zcount = 0
            with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                for it in items:
                    if not it.get("video_id"):
                        continue
                    b = get_thumb_bytes(it["video_id"])
                    if not b:
                        continue
                    prefix = format_number(it["number"], num_fmt) if auto_num else ""
                    base = safe_name(it.get("my_title") or it.get("competitor_title") or it["video_id"])
                    zf.writestr(f"{prefix}{base}.jpg", b)
                    zcount += 1
            zip_buf.seek(0)
            if zcount:
                st.download_button(
                    "Download ZIP (all thumbs)",
                    data=zip_buf.getvalue(),
                    file_name="thumbnails.zip",
                    mime="application/zip",
                    use_container_width=True,
                    key="dl_zip_browser",
                )
            else:
                st.caption("No thumbs available for ZIP")
            if not is_cloud_host():
                if st.button("Save All Thumbs to Folder", use_container_width=True, key="save_all_thumbs"):
                    if require_save_folder():
                        with st.spinner("Saving all thumbnails..."):
                            saved = []
                            for it in items:
                                if it.get("video_id"):
                                    path = ensure_local_thumb(it, num_fmt, auto_num)
                                    if path:
                                        saved.append(path)
                            st.session_state.last_saved = saved
                        st.success(f"✓ {len(saved)} JPG(s) → `{st.session_state.output_dir}`")

        with e3:
            st.markdown("**Combined / PDF**")
            # PDF bytes via temp — browser download, no leftover JPGs
            if st.button("Prepare PDF", use_container_width=True, key="prep_pdf"):
                with st.spinner("Building PDF..."):
                    with tempfile.TemporaryDirectory(prefix="yt_pdf_") as td:
                        pdf_tmp = Path(td) / "export.pdf"
                        make_pdf(items, str(pdf_tmp), num_fmt)
                        st.session_state["pdf_bytes"] = pdf_tmp.read_bytes()
                        st.session_state["pdf_name"] = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                st.success("PDF ready — use Download PDF below.")
            if st.session_state.get("pdf_bytes"):
                st.download_button(
                    "Download PDF",
                    data=st.session_state["pdf_bytes"],
                    file_name=st.session_state.get("pdf_name", "export.pdf"),
                    mime="application/pdf",
                    use_container_width=True,
                    key="dl_pdf_browser",
                )
            if not is_cloud_host():
                if st.button("Save PDF to Folder", use_container_width=True, key="gen_pdf"):
                    if require_save_folder():
                        with st.spinner("Creating PDF..."):
                            pdf_path = Path(st.session_state.output_dir) / f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                            make_pdf(items, str(pdf_path), num_fmt)
                            st.session_state.last_saved = [str(pdf_path)]
                        st.success(f"✓ PDF saved:\\n`{pdf_path}`")


st.markdown("---")
st.caption("Free • Long-form YouTube only • Hosted on Streamlit")
